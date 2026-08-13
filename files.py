import os
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document
from pathlib import Path


def pdf_extracter(path_or_bytes):
    try:
        if isinstance(path_or_bytes, (bytes, bytearray)):
            reader = PdfReader(BytesIO(path_or_bytes))
        else:
            reader = PdfReader(path_or_bytes)

        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error reading PDF {path_or_bytes}: {e}")
        return ""


def docx_extracter(path_or_bytes):
    try:
        if isinstance(path_or_bytes, (bytes, bytearray)):
            doc = Document(BytesIO(path_or_bytes))
        else:
            doc = Document(path_or_bytes)

        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return "\n".join(text).strip()
    except Exception as e:
        print(f"Error reading DOCX {path_or_bytes}: {e}")
        return ""


def extract_text_from_bytes(filename, file_bytes):
    extension = Path(filename).suffix.lower()
    if extension == ".pdf":
        return pdf_extracter(file_bytes)
    if extension == ".docx":
        return docx_extracter(file_bytes)
    if extension == ".txt":
        return file_bytes.decode("utf-8", errors="ignore").strip()
    raise ValueError(f"Unsupported file type: {extension}")


def Text_extracter(path=None):
    if not path:
        return ""

    candidate = Path(path)
    if candidate.suffix.lower() == ".pdf":
        return pdf_extracter(candidate)
    if candidate.suffix.lower() == ".docx":
        return docx_extracter(candidate)

    return ""
